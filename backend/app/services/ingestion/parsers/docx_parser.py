from __future__ import annotations

import io
import docx
from app.services.ingestion.parsers.base_parser import (
    BaseParser,
    ParseResult,
    ParseStatus,
    ParserType,
)


class DocxParser(BaseParser):
    """Parses Word (.docx) documents, extracting paragraphs, tables, and properties metadata."""

    async def parse(
        self,
        *,
        filename: str,
        content_type: str,
        data: bytes,
    ) -> ParseResult:
        try:
            doc = docx.Document(io.BytesIO(data))
            
            # Paragraph extraction
            paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
            
            # Table extraction - formatting cells with pipe delimiters
            tables_text = []
            for table in doc.tables:
                for row in table.rows:
                    row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
                    if row_text:
                        tables_text.append(row_text)
            
            all_text_elements = paragraphs + tables_text
            extracted_text = "\n\n".join(all_text_elements)
            
            status = ParseStatus.SUCCESS if extracted_text.strip() else ParseStatus.EMPTY
            
            # Metadata extraction
            core_properties = doc.core_properties
            metadata = {
                "content_type": content_type,
                "file_name": filename,
                "parser_type": ParserType.DOCX.value,
                "extraction_status": status.value,
                "title": core_properties.title or None,
                "author": core_properties.author or None,
                "created": str(core_properties.created) if core_properties.created else None,
            }
            
            return ParseResult(
                parser_type=ParserType.DOCX,
                file_name=filename,
                status=status,
                extracted_text=extracted_text,
                metadata=metadata,
            )
        except Exception as exc:
            return ParseResult(
                parser_type=ParserType.DOCX,
                file_name=filename,
                status=ParseStatus.FAILED,
                extracted_text="",
                metadata={
                    "content_type": content_type,
                    "error": str(exc),
                    "parser_type": ParserType.DOCX.value,
                    "extraction_status": ParseStatus.FAILED.value,
                },
            )
