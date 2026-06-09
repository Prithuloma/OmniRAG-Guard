"""
tests/test_parsers.py
---------------------
Unit tests for parser contracts and placeholder implementations.
"""

from __future__ import annotations

import fitz
import pytest

from app.services.ingestion.parser_dispatcher import ParserDispatcher
from app.services.ingestion.parsers.base_parser import ParseStatus, ParserType
from app.services.ingestion.parsers.docx_parser import DocxParser
from app.services.ingestion.parsers.image_parser import ImageParser
from app.services.ingestion.parsers.pdf_parser import PDFParser
from app.services.ingestion.parsers.text_parser import TextParser
from unittest.mock import MagicMock, patch
import docx
import io


def _make_pdf_with_text(text: str) -> bytes:
    document = fitz.open()
    page = document.new_page()
    page.insert_text((72, 72), text)
    pdf_bytes = document.tobytes()
    document.close()
    return pdf_bytes


def _make_empty_pdf() -> bytes:
    document = fitz.open()
    document.new_page()
    pdf_bytes = document.tobytes()
    document.close()
    return pdf_bytes


@pytest.mark.parametrize(
    ("filename", "content_type", "expected_parser_cls", "expected_parser_type"),
    [
        ("annual_report.pdf", "application/pdf", PDFParser, ParserType.PDF),
        ("image.png", "image/png", ImageParser, ParserType.IMAGE),
        ("notes.txt", "text/plain", TextParser, ParserType.TEXT),
        (
            "report.docx",
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            DocxParser,
            ParserType.DOCX,
        ),
    ],
)
@pytest.mark.asyncio
async def test_dispatcher_returns_correct_parser_instance(
    filename: str,
    content_type: str,
    expected_parser_cls: type,
    expected_parser_type: ParserType,
) -> None:
    dispatcher = ParserDispatcher()
    dispatch_result = await dispatcher.dispatch(
        filename=filename,
        content_type=content_type,
        data=b"",
    )

    assert isinstance(dispatch_result.parser, expected_parser_cls)
    assert dispatch_result.parser_type is expected_parser_type


@pytest.mark.asyncio
async def test_text_parser_success() -> None:
    parser = TextParser()
    result = await parser.parse(
        filename="test.txt",
        content_type="text/plain",
        data=b"Hello World!",
    )
    assert result.status is ParseStatus.SUCCESS
    assert result.extracted_text == "Hello World!"
    assert result.parser_type is ParserType.TEXT


@pytest.mark.asyncio
async def test_text_parser_empty() -> None:
    parser = TextParser()
    result = await parser.parse(
        filename="test.txt",
        content_type="text/plain",
        data=b"  \n  ",
    )
    assert result.status is ParseStatus.EMPTY
    assert result.extracted_text == "  \n  "


@pytest.mark.asyncio
@patch("app.services.ingestion.parsers.image_parser.Image.open")
@patch("app.services.ingestion.parsers.image_parser.pytesseract.image_to_string")
async def test_image_parser_ocr_success(mock_ocr, mock_open) -> None:
    mock_img = MagicMock()
    mock_img.format = "PNG"
    mock_img.size = (100, 100)
    mock_open.return_value = mock_img
    mock_ocr.return_value = "Extracted OCR text"

    parser = ImageParser()
    result = await parser.parse(
        filename="image.png",
        content_type="image/png",
        data=b"fake-image-bytes",
    )
    assert result.status is ParseStatus.SUCCESS
    assert result.extracted_text == "Extracted OCR text"
    assert result.parser_type is ParserType.IMAGE
    assert result.metadata["image_format"] == "PNG"
    assert result.metadata["image_size"] == (100, 100)


@pytest.mark.asyncio
@patch("app.services.ingestion.parsers.image_parser.Image.open")
@patch("app.services.ingestion.parsers.image_parser.pytesseract.image_to_string")
async def test_image_parser_ocr_empty(mock_ocr, mock_open) -> None:
    mock_img = MagicMock()
    mock_open.return_value = mock_img
    mock_ocr.return_value = "   "

    parser = ImageParser()
    result = await parser.parse(
        filename="image.png",
        content_type="image/png",
        data=b"fake-image-bytes",
    )
    assert result.status is ParseStatus.EMPTY
    assert result.extracted_text == "   "


@pytest.mark.asyncio
@patch("app.services.ingestion.parsers.image_parser.Image.open")
@patch("app.services.ingestion.parsers.image_parser.pytesseract.image_to_string")
async def test_image_parser_ocr_fallback(mock_ocr, mock_open) -> None:
    import pytesseract
    mock_img = MagicMock()
    mock_open.return_value = mock_img
    mock_ocr.side_effect = pytesseract.TesseractNotFoundError()

    parser = ImageParser()
    result = await parser.parse(
        filename="image.png",
        content_type="image/png",
        data=b"fake-image-bytes",
    )
    assert result.status is ParseStatus.PLACEHOLDER
    assert "OCR Fallback" in result.extracted_text


@pytest.mark.asyncio
@patch("app.services.ingestion.parsers.image_parser.Image.open")
async def test_image_parser_failure(mock_open) -> None:
    mock_open.side_effect = Exception("failed to load image")

    parser = ImageParser()
    result = await parser.parse(
        filename="image.png",
        content_type="image/png",
        data=b"corrupted-bytes",
    )
    assert result.status is ParseStatus.FAILED
    assert result.extracted_text == ""
    assert "failed to load image" in result.metadata["error"]


def _make_docx(text: str, table_data: list[list[str]] | None = None) -> bytes:
    doc = docx.Document()
    doc.add_paragraph(text)
    if table_data:
        table = doc.add_table(rows=len(table_data), cols=len(table_data[0]))
        for r_idx, row in enumerate(table_data):
            for c_idx, val in enumerate(row):
                table.cell(r_idx, c_idx).text = val
    doc.core_properties.title = "Test Title"
    doc.core_properties.author = "Test Author"
    stream = io.BytesIO()
    doc.save(stream)
    return stream.getvalue()


@pytest.mark.asyncio
async def test_docx_parser_success() -> None:
    docx_bytes = _make_docx("Hello Paragraph", [["Cell1", "Cell2"]])
    parser = DocxParser()
    result = await parser.parse(
        filename="document.docx",
        content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        data=docx_bytes,
    )
    assert result.status is ParseStatus.SUCCESS
    assert "Hello Paragraph" in result.extracted_text
    assert "Cell1 | Cell2" in result.extracted_text
    assert result.metadata["title"] == "Test Title"
    assert result.metadata["author"] == "Test Author"
    assert result.parser_type is ParserType.DOCX


@pytest.mark.asyncio
async def test_docx_parser_empty() -> None:
    docx_bytes = _make_docx("   ")
    parser = DocxParser()
    result = await parser.parse(
        filename="empty.docx",
        content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        data=docx_bytes,
    )
    assert result.status is ParseStatus.EMPTY
    assert result.extracted_text == ""


@pytest.mark.asyncio
async def test_docx_parser_corrupted() -> None:
    parser = DocxParser()
    result = await parser.parse(
        filename="corrupted.docx",
        content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        data=b"invalid-docx-bytes",
    )
    assert result.status is ParseStatus.FAILED
    assert result.extracted_text == ""


@pytest.mark.asyncio
async def test_pdf_parser_extracts_text_from_valid_pdf() -> None:
    parser = PDFParser()
    expected_text = "OmniRAG quarterly report"
    pdf_bytes = _make_pdf_with_text(expected_text)

    result = await parser.parse(
        filename="annual_report.pdf",
        content_type="application/pdf",
        data=pdf_bytes,
    )

    assert result.parser_type is ParserType.PDF
    assert result.file_name == "annual_report.pdf"
    assert result.status is ParseStatus.SUCCESS
    assert expected_text in result.extracted_text
    assert result.metadata["page_count"] == 1
    assert result.metadata["file_name"] == "annual_report.pdf"
    assert result.metadata["parser_type"] == ParserType.PDF.value
    assert result.metadata["extraction_status"] == ParseStatus.SUCCESS.value


@pytest.mark.asyncio
async def test_pdf_parser_handles_empty_pdf() -> None:
    parser = PDFParser()
    pdf_bytes = _make_empty_pdf()

    result = await parser.parse(
        filename="blank.pdf",
        content_type="application/pdf",
        data=pdf_bytes,
    )

    assert result.parser_type is ParserType.PDF
    assert result.status is ParseStatus.EMPTY
    assert result.extracted_text == ""
    assert result.metadata["page_count"] == 1
    assert result.metadata["extraction_status"] == ParseStatus.EMPTY.value


@pytest.mark.asyncio
async def test_pdf_parser_handles_corrupted_pdf() -> None:
    parser = PDFParser()

    result = await parser.parse(
        filename="corrupted.pdf",
        content_type="application/pdf",
        data=b"%PDF-1.4\nthis is not a valid pdf document",
    )

    assert result.parser_type is ParserType.PDF
    assert result.status is ParseStatus.FAILED
    assert result.extracted_text == ""
    assert result.metadata["page_count"] == 0
    assert result.metadata["extraction_status"] == ParseStatus.FAILED.value
